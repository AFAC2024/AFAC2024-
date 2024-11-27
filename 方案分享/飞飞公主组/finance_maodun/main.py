import os
import json
import re
import argparse
from tqdm import tqdm
from docx import Document
from llm import LLM

from template import INSTRUCTION1, INSTRUCTION2, INSTRUCTION3
class DocumentProcessor:
    

    def __init__(self, model_path):
        self.llm = LLM(model_path)
        self.INSTRUCTION1 = INSTRUCTION1
        self.INSTRUCTION2 = INSTRUCTION2
        self.INSTRUCTION3 = INSTRUCTION3

    def process_documents(self, category, process_type):
        result = []
        for document_path in category:
            print(f"正在处理文档:{os.path.basename(document_path)}")
            answer_list = []

            text = ""
            doc = Document(document_path)
            for p in doc.paragraphs:
                text += "\n" + p.text

            if process_type == "bid":
                paragraph_list = self.split_bid_paragraphs(text)
                instruction = self.INSTRUCTION1
            elif process_type == "clause":
                paragraph_list = self.split_clause_paragraphs(text)
                instruction = self.INSTRUCTION2
            elif process_type == "report":
                error_sentence1 = self.find_invalid_Q(text)
                answer_list.extend(error_sentence1)
                paragraph_list = self.split_report_paragraphs(text)
                instruction = self.INSTRUCTION3
            else:
                raise ValueError("Invalid process type")

            for paragraph in tqdm(paragraph_list):
                prompt = instruction + "\n" + paragraph
                answer = self.llm.predict(prompt)
                if answer != "错误句子：无":
                    if process_type == "clause" and "港澳台" in answer:
                        continue
                    answer_list.append([answer[5:]])

            if process_type == "report":
                final_answer_list = []
                for answer in answer_list:
                    checked_answer = self.check_answer(answer, text)
                    if checked_answer:
                        final_answer_list.append(checked_answer)

                answer_list = self.remove_duplicates(final_answer_list)
                result.append({"id": os.path.splitext(os.path.basename(document_path))[0], "sents": final_answer_list})
            else:
                answer_list = self.remove_duplicates(answer_list)
                result.append({"id": os.path.splitext(os.path.basename(document_path))[0], "sents": answer_list})
        return result

    def split_bid_paragraphs(self, text):
        paragraphs = text.strip().split('\n')
        merged_paragraphs = []
        current_paragraph = []
        if "一、" in text and "二、" in text and "三、" in text and "四、" in text and "五、" in text:
            levels = ["一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "十一、", "十二、"]
        else:
            levels = [f"{i}." for i in range(1, 16)]
        current_level_index = 0

        for paragraph in paragraphs:
            # 检查段落是否以数字加句号开头
            if paragraph.strip().startswith(levels[current_level_index]):
                # 如果当前段落是新的级别，则将之前的段落合并并添加到结果中
                if current_paragraph:
                    merged_paragraphs.append('\n'.join(current_paragraph))
                    current_paragraph = []
                    current_level_index += 1

            current_paragraph.append(paragraph)

        # 添加最后一个段落
        if current_paragraph:
            merged_paragraphs.append('\n'.join(current_paragraph))

        # 把过长的字符串切分
        output_list = []
        for string in merged_paragraphs:
            if len(string) > 1024:
                for i in range(0, len(string), 1024):
                    output_list.append(string[i:i + 1024])
            else:
                output_list.append(string)
        return output_list

    def split_clause_paragraphs(self, text):
        paragraphs = text.strip().split('\n')
        merged_paragraphs = []
        current_paragraph = []

        for paragraph in paragraphs:
            # 检查段落是否以“第”开头并在前10个字符内包含“条”
            if "第" in paragraph[:10] and "条" in paragraph[:10]:
                # 如果当前段落是新的级别，则将之前的段落合并并添加到结果中
                if current_paragraph:
                    merged_paragraphs.append('\n'.join(current_paragraph))
                    current_paragraph = []

            current_paragraph.append(paragraph)

        # 添加最后一个段落
        if current_paragraph:
            merged_paragraphs.append('\n'.join(current_paragraph))

        re_merged_paragraphs = []
        small_paragraph = ""
        for paragraph in merged_paragraphs:
            if len(small_paragraph) + len(paragraph) + 1 < 512:  # +1 for the newline character
                if small_paragraph:
                    small_paragraph += "\n" + paragraph
                else:
                    small_paragraph = paragraph
            else:
                re_merged_paragraphs.append(small_paragraph)
                small_paragraph = paragraph

        if small_paragraph:
            re_merged_paragraphs.append(small_paragraph)

        # 把过长的字符串切分
        output_list = []
        for string in merged_paragraphs:
            if len(string) > 1024:
                for i in range(0, len(string), 1024):
                    output_list.append(string[i:i + 1024])
            else:
                output_list.append(string)
        return output_list

    def split_report_paragraphs(self,text):
        paragraphs = text.strip().split('\n')

        return [line for line in paragraphs if line.strip() != '']

    def find_invalid_Q(self, text):
        # 定义正则表达式，用于匹配中文的逗号、句号和换行符
        pattern = r'[，；。\n]'
        
        # 使用正则表达式拆分字符串
        sentences = re.split(pattern, text)
        
        # 用于存储结果的列表
        result = []
        
        for sentence in sentences:
            # 检查句子中是否包含字母 Q
            if 'Q' in sentence:
                # 查找 Q 后面的字符
                q_index = sentence.index('Q')
                if q_index + 1 < len(sentence) and sentence[q_index + 1].isdigit():
                    # 检查数字范围是否是 1 到 4
                    number = int(sentence[q_index + 1])
                    if not (1 <= number <= 4):
                        # 查找句子在原字符串中的位置
                        start_index = text.find(sentence)
                        if start_index != -1:
                            # 查找紧接着的标点符号
                            end_index = start_index + len(sentence)
                            if end_index < len(text):
                                next_char = text[end_index]
                                result.append([sentence.strip() + next_char])
        
        return result

    def check_answer(self, answer, str2):
        str1 = answer[0]
        # 检查str1是否是str2的子串
        if str1 not in str2:
            return []

        # 定义特定标点符号
        specific_punctuations = "，。；"

        # 检查str1中是否只有一个特定标点符号
        punctuation_count = sum(1 for char in str1 if char in specific_punctuations)
        if punctuation_count == 1:
            return [str1]

        # 如果没有特定标点符号
        if punctuation_count == 0:
            # 检查str1在str2中是否后面紧跟着特定标点符号
            index = str2.index(str1)
            if index + len(str1) < len(str2) and str2[index + len(str1)] in specific_punctuations:
                return [str1 + str2[index + len(str1)]]
            else:
                return []

        # 如果有大于一个的特定标点符号
        if punctuation_count > 1:
            # 按照特定标点符号分割str1
            substrings = re.split(r'[，。；]', str1)
            # 判断哪一个子字符串中有数字
            substrings_with_digits = [substring for substring in substrings if any(char.isdigit() for char in substring)]
            if len(substrings_with_digits) == 1:
                substring_with_digits = substrings_with_digits[0]
                index = str1.index(substring_with_digits)
                if index + len(substring_with_digits) < len(str1) and str1[index + len(substring_with_digits)] in specific_punctuations:
                    return [substring_with_digits + str1[index + len(substring_with_digits)]]
                else:
                    return [substring_with_digits]
            else:
                return []

    def remove_duplicates(self, answer_list):
        """
        去除包含子列表的列表中的重复子列表。
        """
        if len(answer_list) >5:
            return []
        unique_answer_list = []
        for sublist in answer_list:
            if sublist not in unique_answer_list:
                unique_answer_list.append(sublist)
        return unique_answer_list

    @staticmethod
    def save_as_json_lines(data, output_file):
        with open(output_file, 'w', encoding='utf-8') as file:
            for item in data:
                json.dump(item, file, ensure_ascii=False)
                file.write('\n')

def split_category(document_path_list):
    category_1 = []
    category_2 = []
    category_3 = []

    for document_path in document_path_list:

        text = ""
        doc = Document(document_path)
        for p in doc.paragraphs:
            text = text + "\n" + p.text

        if ("招标" in text[:200] or "采购" in document_path or "招标" in document_path or "公告" in document_path) and "法" not in document_path:
            category_1.append(document_path)
        elif "条款" in document_path or "条例" in document_path or "办法" in document_path or "中华人民共和国" in document_path:
            category_2.append(document_path)
        else:
            category_3.append(document_path)
    
    return category_1, category_2, category_3

def main(args):
    # 获取当前工作目录路径
    # current_directory = os.getcwd()
    # print(f"当前工作目录: {current_directory}")
    # model_path = os.path.join(current_directory, "../", "models/qwen2-7b-instruct-finance-maodun")
    model_path = args.model_path
    folder = args.input_folder
    save_path = args.output_file

    document_path_list = [os.path.join(folder, filename) for filename in os.listdir(folder) if filename.endswith('.docx')]
    category_1, category_2, category_3 = split_category(document_path_list)

    processor = DocumentProcessor(model_path)
    result1 = processor.process_documents(category_1, "bid")
    result2 = processor.process_documents(category_2, "clause")
    result3 = processor.process_documents(category_3, "report")

    result = result1 + result2 + result3
    DocumentProcessor.save_as_json_lines(result, save_path)
    print(f"处理完毕！结果保存在{save_path}")

if __name__ == "__main__":

    parser = argparse.ArgumentParser(description="Process documents and save results.")
    parser.add_argument("--model_path", type=str, default="/models/qwen2-7b-instruct-finance-maodun-merged", help="Path to the model")
    parser.add_argument("--input_folder", type=str, default="roundb", help="Path to the folder containing documents")
    parser.add_argument("--output_file", type=str, default="result.json", help="Path to save the output JSONL file")
    args = parser.parse_args()

    main(args)